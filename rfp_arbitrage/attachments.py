"""Pull solicitation documents and turn them into text. The TERMS live in the attachments
(the notice page is a summary), so the clause gate is only as good as this stage.

Handles: PDF (pdfminer.six), DOCX (python-docx), HTML, plain text, ZIP (recursed one level).
Anything else is recorded with an error so the gate can say "unverified" instead of "silent"."""
from __future__ import annotations

import io
import logging
import re
import zipfile
from typing import Iterator

from .config import Settings, settings as _settings
from .models import Opportunity
from .sources.base import Http, strip_html
from .store import Store

MAX_BYTES = 40 * 1024 * 1024
MAX_CHARS = 400_000
MAX_PDF_PAGES = 120
logging.getLogger("pdfminer").setLevel(logging.ERROR)   # "should not allow text extraction" warnings on every DRM'd PDF          # pdfminer is ~0.3 s/page; terms live in the first hundred pages of any RFP


def _kind(url: str, content_type: str, head: bytes) -> str:
    u = url.lower().split("?")[0]
    ct = (content_type or "").lower()
    if head.startswith(b"%PDF") or u.endswith(".pdf") or "pdf" in ct:
        return "pdf"
    if head.startswith(b"PK"):
        if u.endswith(".docx") or "wordprocessingml" in ct:
            return "docx"
        if u.endswith(".xlsx") or "spreadsheetml" in ct:
            return "xlsx"
        if u.endswith(".pptx") or "presentationml" in ct:
            return "pptx"
        return "zip"
    if "html" in ct or head.lstrip()[:1] == b"<":
        return "html"
    if "text/" in ct or u.endswith((".txt", ".md", ".csv")):
        return "text"
    return "binary"


def extract_text(data: bytes, kind: str) -> str:
    if kind == "pdf":
        from pdfminer.high_level import extract_text as pdf_text
        return pdf_text(io.BytesIO(data), maxpages=MAX_PDF_PAGES)[:MAX_CHARS]
    if kind == "docx":
        import docx
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)[:MAX_CHARS]
    if kind == "html":
        return strip_html(data.decode("utf-8", "replace"))[:MAX_CHARS]
    if kind == "text":
        return data.decode("utf-8", "replace")[:MAX_CHARS]
    if kind == "zip":
        out = []
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            for name in z.namelist()[:40]:
                if name.endswith("/"):
                    continue
                inner = z.read(name)
                k = _kind(name, "", inner[:8])
                if k in ("pdf", "docx", "html", "text"):
                    try:
                        out.append(f"### {name}\n" + extract_text(inner, k))
                    except Exception as e:  # noqa: BLE001
                        out.append(f"### {name}\n[extract error: {e}]")
        return "\n\n".join(out)[:MAX_CHARS]
    if kind in ("xlsx", "pptx"):
        # cheap: pull the XML strings; enough for the gate to see clause words if present
        out = []
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            for name in z.namelist():
                if name.endswith(".xml") and ("sharedStrings" in name or "slides/slide" in name):
                    out.append(strip_html(z.read(name).decode("utf-8", "replace")))
        return "\n".join(out)[:MAX_CHARS]
    return ""


_SAM_DESC_QUOTA_HIT = False    # process-wide: once the description endpoint 429s, stop asking until the daily reset


class Fetcher:
    def __init__(self, store: Store, cfg: Settings | None = None):
        self.store = store
        self.cfg = cfg or _settings()
        self.http = Http(self.cfg)

    def urls_for(self, opp: Opportunity) -> list[str]:
        urls = list(opp.attachments)
        if opp.source == "sam_gov":
            from .sources.sam_gov import SamGov
            sam = SamGov(self.cfg, self.http)
            urls = [sam.attachment_url(u) for u in urls]
        elif opp.source in ("merx", "canadabuys", "seao_quebec", "socrata", "bidnet") and opp.url:
            urls = [opp.url] + urls        # the notice page itself carries terms on these portals
        return list(dict.fromkeys(u for u in urls if u.startswith("http")))

    def fetch(self, opp: Opportunity, max_docs: int = 12, refresh: bool = False) -> Iterator[dict]:
        # a SAM notice body is only reachable through its description endpoint
        global _SAM_DESC_QUOTA_HIT
        if opp.source == "sam_gov" and not opp.description and not _SAM_DESC_QUOTA_HIT:
            from .sources.sam_gov import SamGov
            try:
                body = SamGov(self.cfg, self.http).describe(opp)
                if body:
                    self.store.put_document(opp.key, opp.raw.get("description_url", "sam:desc"), "html", body)
                    yield {"url": "sam:desc", "kind": "html", "chars": len(body)}
            except Exception as e:  # noqa: BLE001
                if "429" in str(e):
                    _SAM_DESC_QUOTA_HIT = True      # attachments still download from sam.gov itself
                self.store.put_document(opp.key, "sam:desc", "html", "", str(e))
        for url in self.urls_for(opp)[:max_docs]:
            if not refresh and self.store.has_document(opp.key, url):
                continue
            try:
                import requests
                self.http._throttle()
                r = self.http.s.get(url, timeout=90, stream=True, allow_redirects=True)
                r.raise_for_status()
                data = r.raw.read(MAX_BYTES, decode_content=True)
                kind = _kind(r.url, r.headers.get("Content-Type", ""), data[:8])
                text = extract_text(data, kind)
                text = re.sub(r"[ \t]+", " ", text)
                self.store.put_document(opp.key, url, kind, text, "" if text else "no text extracted")
                yield {"url": url, "kind": kind, "chars": len(text)}
            except Exception as e:  # noqa: BLE001
                self.store.put_document(opp.key, url, "error", "", f"{type(e).__name__}: {e}"[:300])
                yield {"url": url, "kind": "error", "chars": 0, "error": str(e)[:200]}
